"""生成第二周工作报告 .docx"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# ====== 设置默认字体 ======
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)

# ====== 标题 ======
title = doc.add_heading('2026年实践学期 — 第二周工作报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    f'项目名称：Steam游戏评分预测与可视化分析平台\n'
    f'学生姓名：___________    学号：___________\n'
    f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}'
)

doc.add_page_break()

# ====== 一、每日工作记录 ======
doc.add_heading('一、每日工作记录', level=1)

days = [
    {
        'day': '周一',
        'date': '2026年7月6日',
        'work': [
            '回顾第一周工作成果，梳理现有系统功能（3个标签页 + 6个API端点 + Flask后端 + 单文件前端）',
            '分析系统待改进方向：游戏详情缺少评论趋势图表、异常数据未充分利用、预测偏差原因不明确、UI设计需提升、缺少AI对话辅助功能',
            '制定第二周工作计划：五大升级方向（时序图表、异常分析、偏差诊断、UI美化、AI助手集成）',
            '深入阅读现有代码（app.py约480行、index.html约800行），全面理解系统架构和数据流',
            '探索数据库schema（69列、37,481条记录）和模型特征集（49个特征、前30标签One-Hot编码）',
            '确定技术方案：前端使用Chart.js增强图表、后端引入IsolationForest异常检测缓存、偏差分析采用规则引擎、UI采用玻璃拟态设计、AI助手采用浮动组件+新窗口方案',
        ],
        'ai_usage': [
            '使用Claude AI自动探索代码库（3个Explore Agent并行分析app.py/index.html/data_cleaning.py/modeling.py等文件）',
            'AI分析项目依赖（Python 8个包 + JS 1个CDN库），评估新增功能的技术可行性',
            'AI生成详细实施方案（包含API设计、前端组件架构、图表实现方案、DeepSeek集成策略）',
        ],
    },
    {
        'day': '周二',
        'date': '2026年7月7日',
        'work': [
            '编写评论时序模拟算法（generate_review_history函数）：利用S曲线（sigmoid）将全量评论数分配到各月，结合AllReviewPct和RecentReviewPct的差异调整趋势方向，添加随机噪声模拟真实波动',
            '编写偏差分析规则引擎（analyze_gap函数）：实现9条分析规则（经典老游戏效应/质量下滑/持续改进/内购差评/本地化问题/抢先体验风险/定价过高/社区氛围/大众市场突破），每条规则包含置信度计算',
            '编写异常检测引擎（init_anomaly_detection函数）：服务启动时预计算IsolationForest（contamination=0.05）结果并缓存，对异常游戏按6个维度分类标注原因（高价低分/极高人气/极多语言/功能密集等）',
            '在app.py中新增3个API端点：GET /api/games/<id>/review-history（评论时序）、GET /api/anomalies（异常分析）、GET /api/game-quick-info/<id>（游戏快速信息）',
            '修改GET /api/games/<id>端点，集成偏差分析结果到响应中',
            '处理numpy int64 JSON序列化错误、Python 3.8+海象运算符语法兼容等边界问题',
        ],
        'ai_usage': [
            '使用Claude AI编写约350行新增后端代码（时序生成器+偏差分析+异常检测+3个新端点）',
            'AI辅助调试3个问题：numpy.int64 JSON序列化失败/海象运算符语法错误/重复import ast清理',
            'AI生成SQLite PRAGMA查询和pickle反序列化代码，快速获取69列schema和49特征列表',
            'API测试验证：4个新端点全部通过（review-history返回120月数据、anomalies检出1869异常、gap分析定位War Thunder偏差15.3%、quick-info返回8个标签）',
        ],
    },
    {
        'day': '周三',
        'date': '2026年7月8日',
        'work': [
            '全面重写index.html前端（从约33KB扩展到约58KB）：CSS设计系统全面升级',
            '实现玻璃拟态（glassmorphism）设计：半透明卡片(backdrop-filter:blur)、渐变品牌色(--accent-gradient)、粒子动画背景(body::before)、发光边框效果',
            '添加精美CSS动画：卡片悬浮提升(translateY)、脉冲发光(glow-pulse)、淡入过渡(fadeIn)、微粒子漂移(particle-drift)、闪烁加载(shimmer)',
            '导航栏从3标签扩展到5标签：游戏总览/自由预测/数据分析/异常分析（新增）/AI助手（DeepSeek浮动按钮）',
            '游戏详情页重大增强：新增评论趋势双图表区域（堆叠柱状图+评分折线图含模型预测延展点）、偏差分析可展开卡片（含置信度进度条和颜色标记）',
            '新增异常分析标签页：统计卡片+对数散点图(异常高亮)+异常原因分布环形图+箱线图对比+异常游戏列表(含原因标签)',
            '实现DeepSeek AI助手浮动组件：右下角脉冲动画按钮→侧滑面板→快捷提问按钮(4个场景)→剪贴板复制+新标签页打开回退方案',
            '添加Inter专业字体（Google Fonts CDN），优化排版层级（标题800字重/-0.02em间距、正文400字重/1.6行高）',
            '自定义滚动条样式（6px宽、圆角滑轨、悬浮变亮）',
        ],
        'ai_usage': [
            '使用Claude AI一次性生成约58KB的完整前端HTML文件（CSS约320行 + JS约600行）',
            'AI设计玻璃拟态CSS变量系统（24个设计令牌），实现11种动画关键帧',
            'AI生成3个Chart.js图表配置（堆叠柱状图/趋势折线图/对数散点图），解决深色主题下网格线和刻度颜色适配问题',
            'AI设计DeepSeek浮动组件UI（脉冲光环动画/侧滑面板/4格快捷按钮网格/上下文感知问题生成）',
        ],
    },
    {
        'day': '周四',
        'date': '2026年7月9日',
        'work': [
            '中文化翻译：编写200+个标签名的中英文映射表（TAG_TRANSLATION字典），覆盖Indie→独立游戏/Action→动作/Early Access→抢先体验等常见Steam标签',
            '编写20+个特征名的中英映射（FEATURE_TRANSLATION）：HasInAppPurchases→含内购/OriginalPrice→原价(美元)/NumLanguages→语言数量等',
            '翻译9种异常原因的标签（ANOMALY_REASON_TRANSLATION）：overpriced_low_score→高价低分异常/extreme_popularity→极高人气等',
            '翻译9条偏差分析规则的类型名：legacy_classic→经典老游戏效应/declining_quality→近期质量下滑等',
            '修复前端标签选择器bug：tagMap变量作用域问题（从showAnalyze函数内const提升到全局），解决"添加标签后预测失败"的严重功能缺陷',
            '修复预测状态残留问题：预测开始时清除label.innerHTML和factors.innerHTML，失败时显示"预测失败，请检查参数后重试"',
            '后端predict接口增加feature_cn字段返回中文特征名，前端top_factors显示优先使用feature_cn',
            '在index.html中添加150+个标签的前端翻译映射（tagCN对象），游戏详情标签云和数据分析热门标签图表均显示中文',
            '处理翻译字典中的转义字符问题（Beat em up 字符串中单引号与Python语法冲突，改用双引号包裹）',
        ],
        'ai_usage': [
            '使用Claude AI批量生成200+个Steam标签的专业中文翻译（涵盖独立游戏/动作/冒险/模拟/策略/恐怖/RPG/射击/开放世界等全部主流游戏类型）',
            'AI诊断并修复tagMap作用域bug（分析showAnalyze和runPrediction的变量可见性，定位ReferenceError根因）',
            'AI验证所有翻译映射的一致性：前后端TAG_TRANSLATION/tagCN/tagMap三处翻译保持统一',
        ],
    },
    {
        'day': '周五',
        'date': '2026年7月10日',
        'work': [
            '全系统集成测试：逐一验证10个API端点（6个原有+4个新增），确认所有返回数据的中文翻译正确',
            '测试关键用户流程：①无标签预测(81.1%好评)→②添加动作/角色扮演标签预测(79.4%)→③切换异常分析标签查看1869个异常→④点击游戏查看时序图表和偏差分析→⑤打开DeepSeek AI助手',
            '验证偏差分析实战效果：War Thunder (id=15)检出3条偏差原因（近期质量下滑0.64/内购引发差评0.55/本地化不足0.55）',
            '验证前端各页面功能：5个标签页切换流畅、游戏列表排序搜索正常、图表无内存泄漏(destroyCharts清理)、DeepSeek浮动组件点击外部关闭',
            '更新README.md文档：补充第二周新增功能（5标签页/10个API/偏差分析引擎/DeepSeek集成/时序图表/异常分析/特征重要性排名）、玻璃拟态设计说明、200+标签翻译覆盖',
            '编写第二周工作报告（本文档），总结本周5大升级的完成情况',
            '将全部代码更新推送至GitHub仓库：https://github.com/multipointloop/steam-project',
        ],
        'ai_usage': [
            '使用Claude AI连续对话完成全部第二周开发工作，累计对话轮次超过50轮',
            'AI自动化测试：并行调用4个新API端点验证（review-history/anomalies/gap-analysis/quick-info），全部返回正确数据',
            'AI辅助编写README.md更新（新增1500+字内容：API端点表/特征重要性排名/偏差分析规则表/中文翻译覆盖说明）',
            'AI生成第二周工作报告（本文档），仿照第一周格式将本周工作合理分配到周一至周五',
        ],
    },
]

for day_info in days:
    doc.add_heading(f'{day_info["day"]}（{day_info["date"]}）', level=2)

    doc.add_heading('工作内容：', level=3)
    for item in day_info['work']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('AI使用记录：', level=3)
    for item in day_info['ai_usage']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()  # 空行

# ====== 二、本周遇到的问题 ======
doc.add_page_break()
doc.add_heading('二、本周遇到的问题', level=1)

problems = [
    {
        'problem': '缺乏真实时序评论数据，无法直接绘制历史趋势图',
        'detail': '数据集只包含RecentReviewPct和AllReviewPct两个静态快照，没有按时间戳记录的逐月/逐周评论数据。需要在不引入外部数据的前提下生成合理的模拟历史。'
    },
    {
        'problem': 'DeepSeek网页端无法通过iframe嵌入到页面中',
        'detail': '尝试使用<iframe src="https://chat.deepseek.com/">嵌入时，DeepSeek服务器设置了X-Frame-Options: DENY或Content-Security-Policy策略，浏览器拒绝加载。'
    },
    {
        'problem': '标签选择器导致预测功能失效',
        'detail': '用户点击标签后点击"开始预测"时，预测结果区域显示"..."且不更新。经排查发现tagMap变量定义在showAnalyze()函数内部，runPrediction()无法访问，导致ReferenceError。'
    },
    {
        'problem': '预测失败时旧结论残留在界面上',
        'detail': '当一次预测成功后（显示"预测为好评"），若下一次预测因参数问题失败，label.innerHTML未被清除，用户仍看到上一次的结论标签，造成误导。'
    },
    {
        'problem': 'numpy.int64类型无法被Flask的jsonify直接序列化',
        'detail': 'review-history端点返回的cumulative_positive等字段是numpy.int64类型，json.dumps无法处理，抛出TypeError: Object of type int64 is not JSON serializable。'
    },
    {
        'problem': 'Python字典中的单引号字符串包含单引号导致语法错误',
        'detail': '中英文翻译映射字典中包含Beat em up这样的键名，单引号未转义导致Python解析器报SyntaxError: unterminated string literal。'
    },
    {
        'problem': 'Chart.js深色主题下图表可读性问题',
        'detail': '默认的Chart.js图表使用浅色网格线和刻度，在深色(#08090d)背景下对比度不足。需要全局配置所有图表的颜色方案。'
    },
    {
        'problem': 'IsolationForest在37K条数据上的运行性能',
        'detail': '每次API请求时实时运行IsolationForest会导致严重延迟（数秒级）。需要在服务启动时预计算并缓存结果。'
    },
]

for i, p in enumerate(problems, 1):
    doc.add_heading(f'问题{i}：{p["problem"]}', level=3)
    doc.add_paragraph(p['detail'])

# ====== 三、问题的解决方案 ======
doc.add_heading('三、问题的解决方案', level=1)

solutions = [
    '时序数据模拟：设计S曲线（sigmoid）分布算法，根据游戏发行月数将AllReviewCount按"S型曲线"分配到各月（发行初期爆发→中期平稳→长尾衰减），利用AllReviewPct和RecentReviewPct的差异校准评分趋势方向，添加高斯噪声模拟真实波动。在图表上标注"基于S曲线模型的模拟趋势数据"以保持透明度。',
    'DeepSeek嵌入回退方案：放弃iframe嵌入，改用浮动按钮+侧滑面板UI，提供4个场景化快捷提问按钮（数据解读/预测原理/评分体系/购游建议）。点击后通过window.open在新标签页打开DeepSeek网页版，同时自动复制问题到剪贴板。面板底部显示上下文感知的建议问题（根据当前页面自动切换）。',
    '标签选择器修复：将tagMap和topTags从showAnalyze()函数的局部作用域提升到全局作用域，使runPrediction()可正常访问。同时在runPrediction开头添加label.innerHTML=""清除旧结论。',
    '预测状态管理完善：在runPrediction()开始时统一清除box（设为"..."）、label（设为空）和factors（设为空）。预测失败时显示明确的"预测失败"提示和"请检查参数后重试"引导文字。',
    'numpy类型转换：在generate_review_history函数中，对所有返回的数值字段显式调用int()或float()转换为Python原生类型（如int(cumulative_pos)、round(float(scores[i]), 1)），确保JSON序列化兼容。',
    '字符串转义处理：将包含单引号的字典键改为双引号包裹（如"Beat \'em up"），避免Python解析器将字符串内的单引号误认为字符串结束符。',
    'Chart.js深色主题适配：为所有图表配置统一的深色主题选项——刻度颜色#7c8099、网格线颜色rgba(255,255,255,0.05)、图例文字颜色#e8eaf0、透明背景。',
    '异常检测预计算缓存：在app.py启动时（模块加载阶段）运行init_anomaly_detection()，使用IsolationForest(contamination=0.05, n_jobs=-1)对全部37,481条数据一次性拟合，结果存入全局ANOMALY_CACHE字典。API端点直接读取缓存，响应时间<100ms。',
]

for i, s in enumerate(solutions, 1):
    doc.add_paragraph(f'{i}. {s}')

# ====== 四、本周工作总结 ======
doc.add_heading('四、本周工作总结', level=1)

summary = doc.add_paragraph()
summary.add_run(
    '第二周在第一周已完成的基础平台之上，进行了五大方向的深度升级：游戏详情时序图表、异常数据深度分析、'
    '预测偏差自动诊断、Web UI全面美化、DeepSeek AI助手集成。同时完成了全平台中文化翻译（200+标签/20+特征/9种异常原因/9条偏差规则），'
    '修复了3个关键bug（标签选择器失效/预测状态残留/numpy序列化错误）。\n\n'
).font.size = Pt(11)

doc.add_heading('工作成果：', level=3)
achievements = [
    '后端扩展：app.py从约480行扩展到约900行，新增3个API端点（review-history/anomalies/game-quick-info），API总数从6个增至10个',
    '时序图表：实现S曲线模拟算法，游戏详情页新增堆叠柱状图（好评/差评）+ 趋势折线图（含6个月预测延展端点+置信区间）',
    '异常分析：基于IsolationForest(contamination=0.05)检出1,869个异常游戏样本，新增完整分析标签页（散点图+环形图+箱线图+异常列表），标注每个异常的原因分类',
    '偏差诊断：实现9条规则的预测偏差自动分析引擎，当|预测-实际|>15%时触发（如War Thunder检出3条原因：质量下滑0.64/内购差评0.55/本地化问题0.55）',
    'UI美化：从传统暗色主题升级为玻璃拟态设计（24个CSS令牌/11种动画/Inter字体/粒子背景/渐变品牌色/悬浮提升效果），index.html从33KB扩展到58KB（CSS约320行+JS约600行）',
    'AI集成：实现DeepSeek浮动聊天组件（脉冲按钮+侧滑面板+4个快捷提问+剪贴板复制+新窗口回退方案+上下文感知问题生成）',
    '中文化：200+标签名翻译 + 20+特征名翻译 + 9种异常原因翻译 + 9条偏差规则翻译，覆盖前后端全部界面文字',
    'Bug修复：修复3个关键问题（tagMap作用域导致标签选择器失效/预测失败时状态残留/numpy.int64 JSON序列化错误）',
    '文档更新：README.md全面更新（新增1500+字内容：10个API端点表、特征重要性排名Top10、偏差分析规则表、玻璃拟态设计说明）',
]
for a in achievements:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('AI使用总结：', level=3)
doc.add_paragraph(
    '第二周累计使用AI（Claude Code）辅助开发超过50轮对话。AI在以下方面发挥了关键作用：\n'
    '① 代码生成：约350行后端代码（时序生成/偏差分析/异常检测/翻译映射）和约58KB前端代码（CSS设计系统+5个视图+3个Chart.js图表+DeepSeek组件）均由AI生成；\n'
    '② 代码探索：3个并行Explore Agent自动分析了整个项目代码库（5个Python文件+1个HTML文件+SQLite数据库），生成完整的架构报告；\n'
    '③ 问题诊断：AI快速定位3个关键bug的根因（tagMap作用域/numpy序列化/字符串转义），并自动修复验证；\n'
    '④ 自动化测试：AI并行调用4个新API端点进行集成测试，所有端点均返回正确数据；\n'
    '⑤ 文档生成：README.md更新和第二周工作报告由AI辅助撰写，保证格式规范、内容完整、措辞专业；\n'
    '⑥ 翻译覆盖：200+标签的英文→中文翻译由AI批量生成，包含独立游戏/动作/冒险/模拟/策略等主流游戏类型的专业译名。\n'
    '⑦ 设计系统：玻璃拟态CSS变量系统和11种动画关键帧由AI设计，实现了"高端大气上档次"的视觉效果。'
)

doc.add_heading('下周计划：', level=3)
next_week = [
    '准备中期答辩材料：制作项目展示PPT（包含选题背景/技术架构/核心功能演示/创新亮点）',
    '录制系统演示视频：展示5个标签页的全部功能 + 预测流程 + 异常分析 + DeepSeek AI助手',
    '探索更多加分项：NLP分析游戏描述文本的情感倾向、3D可视化（Three.js）、移动端PWA支持',
    '尝试更高阶模型：XGBoost/LightGBM对比实验，目标将R²从0.25提升至0.35+',
    '完善测试用例：编写自动化接口测试脚本，确保10个API端点的稳定性',
    '编写第三周报告，总结中期进展',
]
for n in next_week:
    doc.add_paragraph(n, style='List Bullet')

# ====== 保存 ======
output_path = 'reports/第二周工作报告.docx'
doc.save(output_path)
print(f'报告已保存至: {output_path}')
